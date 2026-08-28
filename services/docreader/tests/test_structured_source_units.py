from __future__ import annotations

from io import BytesIO

import pytest
from docx import Document as DocxDocument
from openpyxl import Workbook
from openpyxl.worksheet.table import Table, TableStyleInfo
from PIL import Image

from docreader.main import DocReaderServicer, _structured_unit_to_proto
from docreader.models.document import (
    DocumentLocator,
    FormImageParent,
    ImageLocator,
    ParagraphImageParent,
    SpreadsheetLocator,
    StructuredSourceUnitKind,
    TableCellImageParent,
)
from docreader.parser.docx_parser import (
    _docx_package_image_payloads,
    _docx_structured_units,
)
from docreader.parser.excel_parser import ExcelParser
from docreader.parser.image_parser import ImageParser
from docreader.parser.pdf_parser import PDFParser, _pdf_image_unit
from docreader.proto.docreader_pb2 import ReadRequest


def _image_bytes(fmt: str) -> bytes:
    output = BytesIO()
    Image.new("RGB", (4, 3), (10, 20, 30)).save(output, format=fmt)
    return output.getvalue()


@pytest.mark.parametrize(
    ("extension", "fmt", "mime"),
    [("png", "PNG", "image/png"), ("jpg", "JPEG", "image/jpeg"), ("webp", "WEBP", "image/webp")],
)
def test_image_parser_emits_typed_region_without_ocr(extension: str, fmt: str, mime: str) -> None:
    document = ImageParser(file_name=f"proof.{extension}", file_type=extension).parse_into_text(
        _image_bytes(fmt)
    )
    assert len(document.structured_source_units) == 1
    unit = document.structured_source_units[0]
    assert unit.key == "image:0"
    assert unit.kind is StructuredSourceUnitKind.IMAGE_REGION
    assert unit.text == ""
    assert isinstance(unit.locator, ImageLocator)
    assert (unit.locator.width, unit.locator.height, unit.locator.media_type) == (4, 3, mime)


def test_compound_image_parent_rejects_mixed_or_unrelated_coordinates() -> None:
    from pydantic import ValidationError

    with pytest.raises(ValidationError):
        ImageLocator.model_validate(
            {
                "original_ref": "images/a.png",
                "width": 1,
                "height": 1,
                "media_type": "image/png",
                "page_ordinal": 0,
                "compound_parent": {
                    "parent_kind": "paragraph",
                    "section_ordinal": 0,
                    "paragraph_ordinal": 0,
                },
            }
        )
    with pytest.raises(ValidationError):
        ImageLocator.model_validate(
            {
                "original_ref": "images/a.png",
                "width": 1,
                "height": 1,
                "media_type": "image/png",
                "compound_parent": {
                    "parent_kind": "paragraph",
                    "section_ordinal": 0,
                    "paragraph_ordinal": 0,
                    "table_ordinal": 9,
                },
            }
        )


def test_xlsx_preserves_sheets_cells_merges_tables_and_rows() -> None:
    workbook = Workbook()
    first = workbook.active
    assert first is not None
    first.title = "Requirements"
    first.append(["Item", "Response"])
    first.append(["Security", "Comply"])
    first.append(["Support", "24x7"])
    first.merge_cells("C6:D6")
    first["C6"] = "Later merged instruction"
    first.merge_cells("A4:B4")
    first["A4"] = "Merged instruction"
    table = Table(displayName="RequirementsTable", ref="A1:B3")
    table.tableStyleInfo = TableStyleInfo(name="TableStyleMedium2")
    first.add_table(table)
    second = workbook.create_sheet("Pricing")
    second["A1"] = "Total"
    second["B1"] = 42
    merge_cases = workbook.create_sheet("MergeCases")
    merge_cases["A1"] = "Vertical"
    merge_cases.merge_cells("A1:A2")
    for address, value in {
        "D1": "Key",
        "E1": "Value",
        "D2": "Contained",
        "D3": "Crossing",
        "E3": "extends",
        "D4": "Row merge",
        "F4": "tail",
    }.items():
        merge_cases[address] = value
    merge_cases.merge_cells("D2:E2")
    merge_cases.merge_cells("E3:F3")
    merge_cases.merge_cells("D4:E4")
    merge_cases.add_table(Table(displayName="MergeTable", ref="D1:E3"))
    output = BytesIO()
    workbook.save(output)

    document = ExcelParser(file_name="tender.xlsx", file_type="xlsx").parse_into_text(
        output.getvalue()
    )
    units = document.structured_source_units
    assert [unit.key for unit in units if unit.kind is StructuredSourceUnitKind.SECTION] == [
        "sheet:0",
        "sheet:1",
        "sheet:2",
    ]
    table_unit = next(unit for unit in units if unit.key == "sheet:0:table:RequirementsTable")
    assert table_unit.kind is StructuredSourceUnitKind.TABLE_REGION
    assert isinstance(table_unit.locator, SpreadsheetLocator)
    assert table_unit.locator.defined_tables[0].a1_range == "A1:B3"
    sheet = next(unit for unit in units if unit.key == "sheet:0")
    assert isinstance(sheet.locator, SpreadsheetLocator)
    assert [region.a1_range for region in sheet.locator.merged_ranges] == [
        "A4:B4",
        "C6:D6",
    ]
    assert [(cell.address, cell.text) for cell in sheet.locator.cells][:2] == [
        ("A1", "Item"),
        ("B1", "Response"),
    ]
    assert any(unit.key == "sheet:1:row:1" and unit.text == "Total | 42" for unit in units)
    merge_sheet = next(unit for unit in units if unit.key == "sheet:2")
    assert isinstance(merge_sheet.locator, SpreadsheetLocator)
    assert [item.a1_range for item in merge_sheet.locator.merged_ranges] == [
        "A1:A2",
        "D2:E2",
        "E3:F3",
        "D4:E4",
    ]
    merge_table = next(unit for unit in units if unit.key == "sheet:2:table:MergeTable")
    assert isinstance(merge_table.locator, SpreadsheetLocator)
    assert [item.a1_range for item in merge_table.locator.merged_ranges] == ["D2:E2"]
    row_four = next(unit for unit in units if unit.key == "sheet:2:row:4")
    assert isinstance(row_four.locator, SpreadsheetLocator)
    assert [item.a1_range for item in row_four.locator.merged_ranges] == ["D4:E4"]
    assert [unit.ordinal for unit in units] == list(range(len(units)))


def test_xlsx_sparse_last_cell_is_rejected_without_rectangular_iteration() -> None:
    from docreader.parser.excel_parser import XlsxStructureLimitError

    workbook = Workbook()
    sheet = workbook.active
    assert sheet is not None
    sheet["XFD1048576"] = "sparse bomb"
    output = BytesIO()
    workbook.save(output)

    with pytest.raises(
        XlsxStructureLimitError,
        match=r"XLSX_STRUCTURE_LIMIT_EXCEEDED:logical_dimension_end_row",
    ):
        ExcelParser(file_name="sparse.xlsx", file_type="xlsx").parse_into_text(
            output.getvalue()
        )


def test_docx_preserves_body_order_heading_owner_and_drawing_identity() -> None:
    document = DocxDocument()
    document.add_heading("Heading A", level=1)
    document.add_paragraph("Narrative")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Requirement"
    table.cell(0, 1).text = "Answer"
    table.cell(1, 0).text = "Encryption"
    table.cell(1, 1).text = "Yes"
    document.add_heading("Heading B", level=1)
    document.add_picture(BytesIO(_image_bytes("PNG")))
    output = BytesIO()
    document.save(output)

    units = _docx_structured_units(output.getvalue())
    keys = [unit.key for unit in units]
    assert keys.index("section:0") < keys.index("table:0") < keys.index("section:1")
    table_unit = next(unit for unit in units if unit.key == "table:0")
    assert table_unit.kind is StructuredSourceUnitKind.TABLE_REGION
    assert isinstance(table_unit.locator, DocumentLocator)
    assert table_unit.locator.section_ordinal == 0
    assert table_unit.locator.heading_path == "Heading A"
    assert any(unit.key == "table:0:row:1" and "Encryption" in unit.text for unit in units)
    image_unit = next(unit for unit in units if unit.kind is StructuredSourceUnitKind.IMAGE_REGION)
    assert isinstance(image_unit.locator, ImageLocator)
    assert isinstance(image_unit.locator.compound_parent, ParagraphImageParent)
    assert image_unit.locator.compound_parent.section_ordinal == 1
    assert image_unit.locator.compound_parent.paragraph_ordinal >= 0
    assert image_unit.locator.original_ref.startswith("images/")
    assert image_unit.locator.original_ref in _docx_package_image_payloads(output.getvalue())
    assert (image_unit.locator.width, image_unit.locator.height) == (4, 3)
    assert [unit.ordinal for unit in units] == list(range(len(units)))
    section_a = next(unit for unit in units if unit.key == "section:0")
    assert section_a.text == "Heading A\nNarrative"
    assert sum(unit.text.count("Narrative") for unit in units) == 1


def test_docx_headingless_narrative_table_narrative_order_and_typed_image_parents() -> None:
    from docx.oxml import OxmlElement

    document = DocxDocument()
    document.add_paragraph("Before table")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Cell image"
    table.cell(0, 1).paragraphs[0].add_run().add_picture(BytesIO(_image_bytes("PNG")))
    document.add_paragraph("After table")

    form_paragraph = document.add_paragraph()
    form_paragraph.add_run().add_picture(BytesIO(_image_bytes("JPEG")))
    form_xml = OxmlElement("w:sdt")
    form_content = OxmlElement("w:sdtContent")
    form_xml.append(form_content)
    form_paragraph._p.getparent().remove(form_paragraph._p)
    form_content.append(form_paragraph._p)
    document.element.body.insert(-1, form_xml)

    output = BytesIO()
    document.save(output)
    units = _docx_structured_units(output.getvalue())
    keys = [unit.key for unit in units]
    assert keys.index("section:0") < keys.index("table:0") < keys.index("section:1") < keys.index("form:0")
    assert next(unit for unit in units if unit.key == "section:0").text == "Before table"
    assert next(unit for unit in units if unit.key == "section:1").text == "After table"
    assert sum(unit.text.count("Before table") for unit in units) == 1
    assert sum(unit.text.count("After table") for unit in units) == 1

    image_units = [unit for unit in units if unit.kind is StructuredSourceUnitKind.IMAGE_REGION]
    assert isinstance(image_units[0].locator, ImageLocator)
    assert isinstance(image_units[0].locator.compound_parent, TableCellImageParent)
    assert image_units[0].locator.compound_parent.table_ordinal == 0
    assert image_units[0].locator.compound_parent.row_ordinal == 0
    assert image_units[0].locator.compound_parent.cell_ordinal == 1
    assert isinstance(image_units[1].locator, ImageLocator)
    assert isinstance(image_units[1].locator.compound_parent, FormImageParent)
    assert image_units[1].locator.compound_parent.form_ordinal == 0
    table_image_proto = _structured_unit_to_proto(image_units[0])
    form_image_proto = _structured_unit_to_proto(image_units[1])
    assert table_image_proto.image.WhichOneof("compound_parent") == "table_cell_parent"
    assert table_image_proto.image.table_cell_parent.cell_ordinal == 1
    assert form_image_proto.image.WhichOneof("compound_parent") == "form_parent"
    assert form_image_proto.image.form_parent.form_ordinal == 0
    assert [unit.ordinal for unit in units] == list(range(len(units)))


def _minimal_pdf() -> bytes:
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 200 200] /Resources << /Font << /F1 5 0 R >> >> /Contents 4 0 R >>",
        b"<< /Length 39 >>\nstream\nBT /F1 12 Tf 20 100 Td (Tender page) Tj ET\nendstream",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    result = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for index, value in enumerate(objects, 1):
        offsets.append(len(result))
        result.extend(f"{index} 0 obj\n".encode())
        result.extend(value + b"\nendobj\n")
    xref = len(result)
    result.extend(f"xref\n0 {len(objects)+1}\n0000000000 65535 f \n".encode())
    for offset in offsets[1:]:
        result.extend(f"{offset:010d} 00000 n \n".encode())
    result.extend(f"trailer << /Size {len(objects)+1} /Root 1 0 R >>\nstartxref\n{xref}\n%%EOF\n".encode())
    return bytes(result)


def test_pdf_preserves_page_structure_where_supported() -> None:
    document = PDFParser(file_name="tender.pdf", file_type="pdf").parse_into_text(_minimal_pdf())
    assert document.metadata["page_count"] == 1
    assert document.structured_source_units
    assert document.structured_source_units[0].key.startswith("page:0:")


def test_pdf_image_unit_preserves_typed_page_bounds() -> None:
    encoded = __import__("base64").b64encode(_image_bytes("PNG")).decode("ascii")
    unit = _pdf_image_unit(
        "page:2:image:0",
        0,
        2,
        "images/figure.png",
        encoded,
        (10.0, 20.0, 110.0, 120.0),
    )
    assert isinstance(unit.locator, ImageLocator)
    assert unit.locator.page_ordinal == 2
    assert (unit.locator.left, unit.locator.top) == (10.0, 20.0)
    assert (unit.locator.right, unit.locator.bottom) == (110.0, 120.0)


def test_empty_xlsx_sheet_remains_a_valid_structured_response() -> None:
    workbook = Workbook()
    output = BytesIO()
    workbook.save(output)
    parsed = ExcelParser(file_name="empty.xlsx", file_type="xlsx").parse_into_text(
        output.getvalue()
    )
    assert parsed.content == ""
    assert parsed.is_valid()
    assert [unit.key for unit in parsed.structured_source_units] == ["sheet:0"]

    service = DocReaderServicer()
    service._parse_request = lambda request: (parsed, "empty.xlsx")  # type: ignore[method-assign]
    request = ReadRequest(file_name="empty.xlsx", file_type="xlsx", request_id="empty")
    inline = service.Read(request, None)
    frames = list(service.ReadStream(request, None))
    assert inline.error == ""
    assert len(inline.structured_source_units) == 1
    assert frames[0].meta.error == ""
    assert len(frames[0].meta.structured_source_units) == 1


def test_inline_and_streaming_meta_have_identical_structured_units() -> None:
    parsed = ImageParser(file_name="proof.png", file_type="png").parse_into_text(
        _image_bytes("PNG")
    )
    service = DocReaderServicer()
    service._parse_request = lambda request: (parsed, "proof.png")  # type: ignore[method-assign]
    request = ReadRequest(file_name="proof.png", file_type="png", request_id="fixed")
    inline = service.Read(request, None)
    frames = list(service.ReadStream(request, None))
    assert frames[0].HasField("meta")
    assert [unit.SerializeToString() for unit in inline.structured_source_units] == [
        unit.SerializeToString() for unit in frames[0].meta.structured_source_units
    ]
    assert inline.structured_source_units[0].image.width == 4
